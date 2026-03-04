"""Xuất đề thi ra PDF và DOCX."""
import io
from backend.schemas.exam import Exam


def _exam_header(exam: Exam) -> str:
    lines = [f"Đề thi: {exam.title}"]
    if exam.exam_code:
        lines.append(f"Mã đề: {exam.exam_code}")
    lines.append(f"Môn: {exam.subject} | Số câu: {len(exam.questions)}")
    return "\n".join(lines)


def export_docx(exam: Exam) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title_p = doc.add_heading(exam.title, level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    if exam.exam_code:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f"Mã đề: {exam.exam_code}")
        run.bold = True
        run.font.size = Pt(12)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Môn: {exam.subject} | Số câu: {len(exam.questions)}")
    doc.add_paragraph()  # spacing

    # Questions
    for i, q in enumerate(exam.questions, 1):
        # Question content
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Câu {i}: ")
        q_run.bold = True
        q_para.add_run(q.content)

        # Options or matching columns
        if q.question_type == "noi_cot" and q.column_a and q.column_b:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Cột A"
            hdr[1].text = "Cột B"
            max_len = max(len(q.column_a), len(q.column_b))
            for j in range(max_len):
                row = table.add_row().cells
                row[0].text = f"{j+1}. {q.column_a[j]}" if j < len(q.column_a) else ""
                row[1].text = f"{chr(97+j)}. {q.column_b[j]}" if j < len(q.column_b) else ""
        elif q.options:
            for opt in q.options:
                doc.add_paragraph(opt, style="List Bullet")

        doc.add_paragraph()  # spacing between questions

    # Answer key section
    doc.add_page_break()
    doc.add_heading("Đáp án", level=2)
    for i, q in enumerate(exam.questions, 1):
        if q.answer:
            doc.add_paragraph(f"Câu {i}: {q.answer}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(exam: Exam) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Load Unicode font
    font_path = _find_font()
    if font_path:
        pdf.add_font("VN", "", font_path, uni=True)
        pdf.add_font("VN", "B", font_path, uni=True)
        font = "VN"
    else:
        font = "Helvetica"

    page_w = pdf.w - pdf.l_margin - pdf.r_margin

    # Title
    pdf.set_font(font, "B", 16)
    pdf.cell(page_w, 10, exam.title, ln=True, align="C")

    # Metadata
    if exam.exam_code:
        pdf.set_font(font, "B", 12)
        pdf.cell(page_w, 8, f"Ma de: {exam.exam_code}", ln=True, align="C")

    pdf.set_font(font, "", 11)
    pdf.cell(page_w, 8, f"Mon: {exam.subject} | So cau: {len(exam.questions)}", ln=True, align="C")
    pdf.ln(5)

    # Questions
    page_w = pdf.w - pdf.l_margin - pdf.r_margin
    for i, q in enumerate(exam.questions, 1):
        pdf.set_x(pdf.l_margin)
        pdf.set_font(font, "B", 11)
        pdf.multi_cell(page_w, 7, f"Cau {i}: {q.content}")

        pdf.set_font(font, "", 11)
        if q.question_type == "noi_cot" and q.column_a and q.column_b:
            col_w = page_w / 2
            max_len = max(len(q.column_a), len(q.column_b))
            for j in range(max_len):
                left = f"  {j+1}. {q.column_a[j]}" if j < len(q.column_a) else ""
                right = f"  {chr(97+j)}. {q.column_b[j]}" if j < len(q.column_b) else ""
                pdf.set_x(pdf.l_margin)
                pdf.cell(col_w, 6, left)
                pdf.cell(col_w, 6, right, ln=True)
        elif q.options:
            for opt in q.options:
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(page_w, 6, f"  {opt}")

        pdf.ln(3)

    # Answer key
    pdf.add_page()
    pdf.set_x(pdf.l_margin)
    pdf.set_font(font, "B", 14)
    pdf.cell(page_w, 10, "Dap an", ln=True)
    pdf.set_font(font, "", 11)
    for i, q in enumerate(exam.questions, 1):
        if q.answer:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(page_w, 7, f"Cau {i}: {q.answer}")

    return bytes(pdf.output())


def _find_font() -> str:
    """Tìm font Unicode hỗ trợ tiếng Việt."""
    import os
    candidates = [
        os.path.expandvars(r"%WINDIR%\Fonts\arial.ttf"),
        os.path.expandvars(r"%WINDIR%\Fonts\tahoma.ttf"),
        os.path.expandvars(r"%WINDIR%\Fonts\segoeui.ttf"),
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path
    return ""
